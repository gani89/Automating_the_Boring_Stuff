import docx

d = docx.Document('word.docx')

d.add_paragraph('Hello this is a paragraph')
d.add_paragraph('Hello this is another paragraph')

d.save('word3.docx')

#print(d.paragraphs[0].text)
#print(d.paragraphs[1].text)

#p = d.paragraphs[1]

#print(p.runs[0].text) # new run when changes in style (bold, italic). Run 0 - is all

#p.runs[0].bold = True
#p.runs[0].underline = True
#p.runs[0].text = 'Italic and underline'

d.save('word2.docx')





